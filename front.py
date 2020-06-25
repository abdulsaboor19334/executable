from tkinter import *
from pyzbar.pyzbar import decode
import cv2
import numpy as np
import threading
import sqlite3
from datetime import datetime
import os 
import tempfile
import re
import docx
import PIL
import qrcode
import time
from tkinter import ttk
from ttkthemes import themed_tk as tk
from PIL import ImageTk,Image




root = tk.ThemedTk()
root.get_themes()
root.set_theme('itft1')
root.title('Al mansoor garments')
root.geometry('1000x600')
root.iconbitmap('icon.ico')
root.config(background='#D6F1F6')
img = ImageTk.PhotoImage(Image.open("logo.png")) 




bottom_label = ttk.Label(root, text='made by: FluffCoders Email: fluffcoding@gmail.com contact# 03359648486, 03129675450',font='times 10 bold', relief=SUNKEN,anchor=W)
bottom_label.pack(side=BOTTOM, fill=X)



top_label = ttk.Label(root, text='\tAl Mansoor Garments', relief=RAISED)
top_label.pack(side=TOP, fill=X)
top_label.config(font=('times', 44))
mainframe = ttk.Frame(root, relief=RAISED, borderwidth=3)
mainframe.pack(side=LEFT, fill=X, ipadx=100, padx=100,pady=50)


def add_stock_scan():
    cap = cv2.VideoCapture(0)
    data = None
    while data is None:

        ret, img = cap.read()
        code = decode(img)
        for bar in code:
            data = str(bar.data.decode('utf-8'))
            bar_entry.insert(INSERT, data)
        cv2.imshow('my add_stock_scan', img)
        if cv2.waitKey(1) == 27:
            break
        if data is not None:
            cv2.destroyAllWindows()

    cap.release()

def sell_stock_scan():
    cap = cv2.VideoCapture(0)
    data = None
    while data is None:

        ret, img = cap.read()
        code = decode(img)
        for bar in code:
            data = str(bar.data.decode('utf-8'))
            sell_barcode_entry.insert(INSERT, data)
        cv2.imshow('my add_stock_scan', img)
        if cv2.waitKey(1) == 27:
            break
        if data is not None:
            cv2.destroyAllWindows()

    cap.release()

def return_stock_scan():
    cap = cv2.VideoCapture(0)
    data = None
    while data is None:

        ret, img = cap.read()
        code = decode(img)
        for bar in code:
            data = str(bar.data.decode('utf-8'))
            if len(data) < 13:
                return_barcode_entry.insert(INSERT, data)
            elif len(data) >= 13:
                detail_barcode_entry.insert(INSERT, data)
        cv2.imshow('my add_stock_scan', img)
        if cv2.waitKey(1) == 27:
            break
        if data is not None:
            cv2.destroyAllWindows()

    cap.release()

def return_scan():
    video = threading.Thread(target=return_stock_scan)
    video.start()


def add_scan():
    video = threading.Thread(target=add_stock_scan)
    video.start()

def sell_scan():
    video = threading.Thread(target=sell_stock_scan)
    video.start()


def sub():
    sale = int(costprice_entry.get()) * float(profit_entry.get())

    conn = sqlite3.connect('garment.db')
    c = conn.cursor()
    c.execute("INSERT into garment(barcode, name, sale_price, cost_price, gender, size, status) VALUES(:bar, :name, :sale, :cost, :gender, :size, :status)", {
        'bar' : bar_entry.get(),
        'name': name_entry.get(),
        'sale': sale,
        'cost' : costprice_entry.get(),
        'gender': gendervar.get(),
        'size': size_entry.get(),
        'status' : statusvar.get()
    })
    bar_entry.delete(0,END)
    name_entry.delete(0,END)
    costprice_entry.delete(0,END)
    size_entry.delete(0,END)
    conn.commit()

def add_stock():
    global bar_entry
    global name_entry
    global saleprice_entry
    global costprice_entry
    global gender_entry
    global size_entry
    global profit_entry
    global statusvar 
    global gendervar 

    statusvar = StringVar()
    gendervar = StringVar()

    gender_c = ['Male', 'Femle']
    status_c = ['inventory', 'sold']

    level = Toplevel(root)
    level.geometry('1000x600')

#  ENTRIES
    bar_entry = ttk.Entry(level, width=50)
    name_entry = ttk.Entry(level, width=50)
    profit_entry = ttk.Entry(level, width=50)
    costprice_entry = ttk.Entry(level, width=50)
    gender_entry = ttk.OptionMenu(level, gendervar, *gender_c)
    size_entry = ttk.Entry(level, width=50)

#  GRID ENTRIES
    bar_entry.grid(row=1, column=1, padx=20)
    name_entry.grid(row=2, column=1, padx=20)
    profit_entry.grid(row=3, column=1,padx=20)
    costprice_entry.grid(row=4, column=1, padx=20)
    gender_entry.grid(row=5, column=1, padx=20, sticky=W+E)
    size_entry.grid(row=6, column=1, padx=20)


#  LABelS
    barcode = ttk.Label(level, text='BarCode')
    name = ttk.Label(level, text='Name')
    profit = ttk.Label(level, text='Profit')
    costprice = ttk.Label(level,text='cost')
    gender = ttk.Label(level,text='gender')
    size = ttk.Label(level,text='size')



    barcode.grid(row=1,column=0)
    name.grid(row=2,column=0)
    profit.grid(row=3,column=0)
    costprice.grid(row=4,column=0)
    gender.grid(row=5,column=0)
    size.grid(row=6,column=0)


    submit = ttk.Button(level,text="SUBMIT", command=sub)
    submit.grid(row=20,column=0, columnspan=2, sticky=W+E)

    but = ttk.Button(level,text='Scan',command=add_scan)
    but.grid(row=21,column=0, columnspan=2, sticky=W+E)

    #  default values 
    profit_entry.insert(0, 1.)
    statusvar.set('inventory')
    gendervar.set('Male')


total = 0
index_item = 1
def sell(event):
    global total
    global price 
    global discountvar
    global index_item
    global del_entry
    text.config(state='normal')
    conn = sqlite3.connect('garment.db')
    c = conn.cursor()
    disc = discountvar.get() / 100 
    discount = 1 - disc
    year = datetime.now().year
    month = datetime.now().month
    day = datetime.now().day
    
    obj = c.execute('SELECT * FROM garment WHERE barcode = :code',{'code' : sell_barcode_entry.get()})
    for x in obj:
        price = x[2] * discount
        total += price
        xx =  str(index_item)+') '+str(x[0]) +'\t| ' + x[1] +'\t| ' + str(price) +'\t| ' + str(discountvar.get())+'%' +'\t| ' + x[-5] +'\t| ' + x[-6] + '\n'
        c.execute("UPDATE garment SET status = 'sold', discount_price=:price, discount= :discount, year = :year, month= :month, day= :day  WHERE barcode = :code", {'code' : sell_barcode_entry.get(),'price':price, 'discount':discountvar.get(), 'month':month,'day':day, 'year':year})
        index_item += 1
        text.insert(INSERT, xx)
    
    conn.commit()

    sell_barcode_entry.delete(0, END)
    text.config(state='disabled')




def sell_stock():
    global sell_barcode_entry
    global text
    global discountvar

    def copy():
        global total
        
        text.config(state="normal")
        text.insert(END,'\n Your total is : '+ str(total))
        pattern = str(datetime.now())
        number = re.findall('\d+',pattern )
        y = ''
        for x in number:
            y += x
        text.insert(4.5, y[-13:-1])
        text.config(state='disabled')
        total = 0
        index_item = 0
        data = text.get(1.0, 'end')
        w = docx.Document()
        w.add_paragraph(str(data))
        qr = qrcode.QRCode(version=1,border=1)
        qr.add_data(y)
        qr.make(fit=True)
        img = qr.make_image(fill='black',back_color="white")
        img.save('qr.png')
        w.add_picture('qr.png',width=docx.shared.Inches(1),height=docx.shared.Inches(1))
        os.startfile('bill.docx','print')
        w.save('bill.docx')

        conn = sqlite3.connect('garment.db')
        c = conn.cursor()

        c.execute("INSERT INTO trans VALUES(:detail, :id)",{'detail': data, 'id':y})

        conn.commit()

    
    discountvar = IntVar()
    dicount_c = [0,5,10,15,20,25,30,35,40,45,50]

    level = Toplevel(root)
    level.geometry('1000x600')


    barcode = ttk.Label(level, text='Barcode')
    barcode.grid(row=0, column=0)

    discount_label = ttk.Label(level, text="Discount")
    discount_label.grid(row=1, column=0)

    sell_barcode_entry = ttk.Entry(level, width=50)
    sell_barcode_entry.grid(row=0, column=1)

    discount = ttk.OptionMenu(level,discountvar,*dicount_c)
    discount.grid(row=1, column=1)
    discountvar.set(0)

    done = ttk.Button(level, text='Done')
    # done.grid(row=1, column=0, columnspan=2, sticky=W+E)
    sell_barcode_entry.bind('<Return>', sell)

    copy = ttk.Button(level, text="Print" , command=copy)
    copy.grid(row=2, column=0, columnspan=2, sticky=W+E)

    scan = ttk.Button(level, text='Scan', command=sell_scan)
    scan.grid(row=3, column=0, columnspan=2, sticky=W+E)

    def dele():
        line = del_entry.get()
        index1 = int(line) + 10.0
        index2 = index1 + 0.9
        text.config(state="normal")
        for _ in range(0,20):
            text.delete(index1,index2)
        text.config(state='disabled')

    frame = ttk.LabelFrame(level,text="delete objects")
    frame.grid(row=100, column=0, columnspan=2, sticky=W+E, padx=20)

    del_entry = ttk.Entry(frame, width=50)
    del_entry.grid(row=4, column=0, columnspan=2, sticky=W+E)


    delete = ttk.Button(frame, text="del" , command=dele)
    delete.grid(row=5, column=0, columnspan=2, sticky=W+E)
    
    f = ttk.Frame(level)
    f.grid(row=0, column=10, rowspan=10)
    scroll = ttk.Scrollbar(f)
    scroll.pack(side=RIGHT,fill=Y)

    text = Text(f, width=60, yscrollcommand=scroll.set)
    text.pack(side=LEFT)

    scroll.config(command=text.yview)

    text.insert(INSERT, '\t\tAl Mansoor Garments \n')
    text.insert(INSERT, '\n\tdate : ' + str(datetime.now()) + '\n')
    text.insert(INSERT,'No : \n')
    text.insert(INSERT,'   =========================================== \n')
    text.insert(INSERT, '\t\t   Recipt \n')
    text.insert(INSERT,"\t* Return only in the next 5 days \n")
    text.insert(INSERT,"\t* No return on iteams with a discount price \n")
    text.insert(INSERT,'\t=============================== \n')
    text.insert(INSERT, 'Code' + '\t| ' + 'Name' + '\t| ' +'Price' + '\t| ' +'Discount' + '\t| ' +'Size' + '\t| ' +'Gender' +'\n' )
    text.config(state="disabled")

def back():
    conn = sqlite3.connect('garment.db')
    c = conn.cursor()
    code = return_barcode_entry.get()
    c.execute("UPDATE garment SET status ='inventory',day = 0, year=0,month=0 WHERE barcode= :code",{'code' : code })
    item = c.execute("SELECT * from garment WHERE barcode=:code",{'code' : code })
    for x in item:
        xx = '\n'+str(x[0]) +'\t| ' + x[1] +'\t| ' + '\t| ' + x[-2] +'\t| ' + x[-3] + '\n'
        text_return.insert(INSERT, 'RETURNED')
        text_return.insert(INSERT, xx)

    
    conn.commit()
    conn.close()
    return_barcode_entry.delete(0, END)




def detail():
    code = detail_barcode_entry.get()
    conn = sqlite3.connect('garment.db')
    c = conn.cursor()
    data = c.execute("SELECT detail FROM trans WHERE trans_id = :code",{'code':code})
    for x in data:
        y = x
    text_return.insert(INSERT, x)
    c.commit()

def return_stock():
    global return_barcode_entry
    global text_return 
    global detail_barcode_entry
    level = Toplevel(root)
    level.geometry('1000x600')

    barcode = ttk.Label(level, text='Product Barcode')
    barcode.grid(row=0, column=0)

    barcode = ttk.Label(level, text='Reciept Barcode')
    barcode.grid(row=1, column=0)

    return_barcode_entry = ttk.Entry(level, width=50)
    return_barcode_entry.grid(row=0, column=1)

    detail_barcode_entry = ttk.Entry(level, width=50)
    detail_barcode_entry.grid(row=1, column=1)

    scan = ttk.Button(level, text='Return', command=back)
    scan.grid(row=3, column=0, columnspan=2, sticky=W+E)

    but = ttk.Button(level, text='Scan', command=return_scan)
    but.grid(row=4, column=0, columnspan=2, sticky=W+E)

    see_detail = ttk.Button(level, text="Detail", command=detail)
    see_detail.grid(row=5, column=0, columnspan=2, sticky=W+E)

    f = ttk.Frame(level)
    f.grid(row=0, column=10, rowspan=10)

    scroll = ttk.Scrollbar(f)
    scroll.pack(side=RIGHT,fill=Y)

    text_return = Text(f, width=60, yscrollcommand=scroll.set)
    text_return.pack(side=LEFT)



    scroll.config(command=text_return.yview)

def detail():
    global var
    global monthvar
    global yearvar
    level = Toplevel(root)
    level.geometry('1000x600')
    var = StringVar()
    var_c= ['inventory', 'sold']
    var.set('sold')

    monthvar = IntVar()
    monthvar_c = [1,2,3,4,5,6,7,8,9,10,11,12]

    yearvar = IntVar()
    yearvar_c = [2020,2021,2022,2023,2024,2025,2026,2027]

    status = ttk.OptionMenu(level, var, *var_c)
    status.grid(row=0,column=1)

    month = ttk.OptionMenu(level,monthvar,*monthvar_c)
    month.grid(row=2,column=1)

    year = ttk.OptionMenu(level,yearvar,*yearvar_c)
    year.grid(row=3,column=1)


    status_label = ttk.Label(level,text="status")
    status_label.grid(row=0,column=0)

    monthvar.set(datetime.now().month)
    yearvar.set(datetime.now().year)

    
    f = ttk.Frame(level)
    f.grid(row=1, column=3)

    scroll = ttk.Scrollbar(f)
    scroll.pack(side=RIGHT,fill=Y)
    

    detail_text = Text(f, width=87, yscrollcommand=scroll.set)
    detail_text.pack(side=LEFT)
    detail_text.config(state='disabled')

    scroll.config(command=detail_text.yview)
    def show_detail():
        global index_item
        conn = sqlite3.connect('garment.db')
        c = conn.cursor()
        if var.get()=='inventory':
            query = c.execute('SELECT * FROM garment WHERE status = :status',{'status': 'inventory'})
            total_saleprice = 0
            total_costprice = 0
            total_discountprice = 0
            for x in query:
                xx =  str(index_item)+') '+str(x[0]) +'|\t' + x[1] +'|\t' + str(x[2]) +'|\t' + str(x[3])+'|\t' + str(x[4]) +'|\t' + str(x[5]) + '%' +'|\t' + str(x[6]) +'|\t' +str(x[7]) +'|\t' +str(x[8]) +'|\t' +str(x[9])+'/'+str(x[10])+'/'+str(x[11])+'\n'        
                total_saleprice += x[2]
                total_costprice += x[3]
                total_discountprice += x[4]
                detail_text.config(state='normal')
                detail_text.insert(INSERT, str(xx))
                detail_text.config(state='disabled')
                index_item += 1

            detail_text.insert(INSERT,'\n\n')
            detail_text.insert(INSERT,'Total sale price:\t'+str(total_saleprice) + '\n')
            detail_text.insert(INSERT,'Total cost price:\t'+str(total_costprice)+ '\n')
            detail_text.insert(INSERT,'Total discount price:\t'+str(total_discountprice)+ '\n')
            detail_text.config(state='disabled')

        elif var.get() =='sold':
            query = c.execute('SELECT * FROM garment WHERE status = :status AND month=:month AND year = :year',{'status': 'sold', 'month':monthvar.get(), 'year':yearvar.get()})
            total_saleprice = 0
            total_costprice = 0
            total_discountprice = 0
            for x in query:
                xx =  str(index_item)+') '+str(x[0]) +'|\t' + x[1] +'|\t' + str(x[2]) +'|\t' + str(x[3])+'|\t' + str(x[4]) +'|\t' + str(x[5]) + '%' +'|\t' + str(x[6]) +'|\t' +str(x[7]) +'|\t' +str(x[8]) +'\n'        
                total_saleprice += x[2]
                total_costprice += x[3]
                total_discountprice += x[4]
                detail_text.config(state='normal')
                detail_text.insert(INSERT, str(xx))
                
                index_item += 1
            
            
            
            detail_text.insert(INSERT,'\n\n')
            detail_text.insert(INSERT,'Total sale price:\t'+str(total_saleprice) + '\n')
            detail_text.insert(INSERT,'Total cost price:\t'+str(total_costprice)+ '\n')
            detail_text.insert(INSERT,'Total discount price:\t'+str(total_discountprice)+ '\n')
            detail_text.config(state='disabled')
        conn.commit()
    def copy():
        data = detail_text.get(1.0, END)
        doc = docx.Document()
        doc.add_paragraph(data)
        doc.save('report.docx')
        os.startfile('report.docx','print')
        index_item = 0
        total_saleprice = 0
        total_costprice = 0
        total_discountprice = 0
                

    print_button = ttk.Button(level, text='Print Report', command=copy)
    print_button.grid(row=0, column=3)

    show_button = ttk.Button(level, text='Show Monthly Detail',command=show_detail)
    show_button.grid(row=0,column=2)

    detail_text.config(state='normal')
    detail_text.insert(INSERT, 'code| name| sale price| cost price| discount price| discount| gender| size| status\n\n')
    detail_text.config(state='disabled')

def daily_detail():
    global detail_var
    global detail_monthvar
    global detail_yearvar
    global detail_dayvar
    level = Toplevel(root)
    level.geometry('1000x600')
    detail_var = StringVar()
    detail_var_c= ['inventory', 'sold']
    detail_var.set('sold')

    detail_monthvar = IntVar()
    detail_monthvar_c = [1,2,3,4,5,6,7,8,9,10,11,12]

    detail_dayvar = IntVar()
    detail_dayvar_c = [1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30,31]

    detail_yearvar = IntVar()
    detail_yearvar_c = [2020,2021,2022,2023,2024,2025,2026,2027]

    detail_status = ttk.OptionMenu(level, detail_var, *detail_var_c)
    detail_status.grid(row=0,column=1)

    detail_month = ttk.OptionMenu(level,detail_monthvar,*detail_monthvar_c)
    detail_month.grid(row=2,column=1)

    detail_year = ttk.OptionMenu(level,detail_yearvar,*detail_yearvar_c)
    detail_year.grid(row=3,column=1)

    detail_day = ttk.OptionMenu(level,detail_dayvar,*detail_dayvar_c)
    detail_day.grid(row=4,column=1)

    detail_status_label = ttk.Label(level,text="status")
    detail_status_label.grid(row=0,column=0)
    
    detail_dayvar.set(datetime.now().day)
    detail_monthvar.set(datetime.now().month)
    detail_yearvar.set(datetime.now().year)

    f = ttk.Frame(level)
    f.grid(row=1, column=3)

    scroll = ttk.Scrollbar(f)
    scroll.pack(side=RIGHT,fill=Y)
    
    detail_text = Text(f, width=87, yscrollcommand=scroll.set)
    detail_text.pack(side=LEFT)
    detail_text.config(state='disabled')

    scroll.config(command=detail_text.yview)


    def show_detail():
        global index_item
        conn = sqlite3.connect('garment.db')
        c = conn.cursor()
        if detail_var.get()=='inventory':
            query = c.execute("SELECT * FROM garment WHERE status = 'inventory' ")
            total_saleprice = 0
            total_costprice = 0
            total_discountprice = 0
            for x in query:
                xx =  str(index_item)+') '+str(x[0]) +'|\t' + x[1] +'|\t' + str(x[2]) +'|\t' + str(x[3])+'|\t' + str(x[4]) +'|\t' + str(x[5]) + '%' +'|\t' + str(x[6]) +'|\t' +str(x[7]) +'|\t' +str(x[8]) +'|\t' +str(x[9])+'/'+str(x[10])+'/'+str(x[11])+'\n'        
                total_saleprice += x[2]
                total_costprice += x[3]
                total_discountprice += x[4]
                detail_text.config(state='normal')
                detail_text.insert(INSERT, str(xx))
                detail_text.config(state='disabled')
                index_item += 1

            detail_text.insert(INSERT,'\n\n')
            detail_text.insert(INSERT,'Total sale price:\t'+str(total_saleprice) + '\n')
            detail_text.insert(INSERT,'Total cost price:\t'+str(total_costprice)+ '\n')
            detail_text.insert(INSERT,'Total discount price:\t'+str(total_discountprice)+ '\n')
            detail_text.config(state='disabled')

        elif detail_var.get() =='sold':
            query = c.execute("SELECT * FROM garment WHERE status = 'sold' AND month=:month AND year = :year AND day = :day",{ 'month':detail_monthvar.get(), 'year':detail_yearvar.get(), 'day':detail_dayvar.get()})
            total_saleprice = 0
            total_costprice = 0
            total_discountprice = 0
            for x in query:
                xx =  str(index_item)+') '+str(x[0]) +'|\t' + x[1] +'|\t' + str(x[2]) +'|\t' + str(x[3])+'|\t' + str(x[4]) +'|\t' + str(x[5]) + '%' +'|\t' + str(x[6]) +'|\t' +str(x[7]) +'|\t' +str(x[8]) +'\n'        
                total_saleprice += x[2]
                total_costprice += x[3]
                total_discountprice += x[4]
                detail_text.config(state='normal')
                detail_text.insert(INSERT, str(xx))
                
                index_item += 1
            
            
            
            detail_text.insert(INSERT,'\n\n')
            detail_text.insert(INSERT,'Total sale price:\t'+str(total_saleprice) + '\n')
            detail_text.insert(INSERT,'Total cost price:\t'+str(total_costprice)+ '\n')
            detail_text.insert(INSERT,'Total discount price:\t'+str(total_discountprice)+ '\n')
            detail_text.config(state='disabled')
        conn.commit()
    def copy():
        data = detail_text.get(1.0, END)
        doc = docx.Document()
        doc.add_paragraph(data)
        doc.save('report.docx')
        os.startfile('report.docx','print')
        index_item = 0

    print_button = ttk.Button(level, text='Print Report', command=copy)
    print_button.grid(row=0, column=3)

    show_button = ttk.Button(level, text='Show daily Detail',command=show_detail)
    show_button.grid(row=0,column=2)

    detail_text.config(state='normal')
    detail_text.insert(INSERT, 'code| name| sale price| cost price| discount price| discount| gender| size| status\n\n')
    detail_text.config(state='disabled')

def make():
    data_username = add_username_entry.get()
    data_password =  add_password_entry.get()
    data_password2 = add_password2_entry.get()

    conn = sqlite3.connect('garment.db')
    c = conn.cursor()
    data = c.execute("SELECT username FROM auth")

    if data_password != data_password2:
        error = Label(add_frame, text='the passwords do not match')
        error.grid(row=5,column=0,columnspan=2)
        raise Exception('error')
    try:
        c.execute("INSERT INTO auth VALUES(:username , :password)",{'username': data_username, 'password': data_password})
        error = ttk.Label(add_frame, text=f'{data_username} created sucessfully')
        error.grid(row=5,column=0,columnspan=2)
    except sqlite3.IntegrityError:
        error = ttk.Label(add_frame, text='this username already exists please try a unique username')
        error.grid(row=5,column=0,columnspan=2)
    c.commit()

def edit():
    old_username = edit_username_entry.get()
    new_password = edit_password_entry.get()

    conn = sqlite3.connect('garment.db')
    c = conn.cursor()
    c.execute("UPDATE auth SET password=:password WHERE username = :username", {'password' : new_password, 'username':old_username})
    conn.commit()

def delete():
    conn = sqlite3.connect('garment.db')
    c = conn.cursor()
    c.execute('DELETE FROM auth WHERE username = :username',{'username':remove_user})
    conn.commit()
def user_management():
    global add_username_entry
    global add_password_entry
    global add_password2_entry
    global add_frame
    global edit_username_entry
    global edit_password_entry
    level = Toplevel(root)
    level.geometry('1000x600')

    add_frame = ttk.LabelFrame(level,text='Add User',relief=SUNKEN)
    add_frame.grid(row=0, column=0)

    edit_frame = ttk.LabelFrame(level,text='Edit User',relief=SUNKEN)
    edit_frame.grid(row=0, column=1)

    delete_frame = ttk.LabelFrame(level,text='Remove user',relief=SUNKEN)
    delete_frame.grid(row=1,column=0)

    add_username_label = ttk.Label(add_frame,text='new username')
    add_password_label = ttk.Label(add_frame,text='new password')
    add_password_label2 = ttk.Label(add_frame,text='confirm password')

    add_username_label.grid(row=0,column=0)
    add_password_label.grid(row=1,column=0)
    add_password_label2.grid(row=2,column=0)

    add_username_entry = ttk.Entry(add_frame,width=50)
    add_password_entry = ttk.Entry(add_frame,width=50,show='*')
    add_password2_entry = ttk.Entry(add_frame,width=50,show='*')

    add_username_entry.grid(row=0,column=1)
    add_password_entry.grid(row=1,column=1)
    add_password2_entry.grid(row=2,column=1)

    create = ttk.Button(add_frame,text='Create user',command=make)
    create.grid(row=3,column=0,columnspan=2)

    edit_button = ttk.Button(edit_frame,text='Edit user',command=edit)
    edit_button.grid(row=4,column=0,columnspan=2)

    edit_username_entry = ttk.Entry(edit_frame,width=50)
    edit_password_entry = ttk.Entry(edit_frame,width=50)
    edit_username_entry.grid(row=0,column=1)
    edit_password_entry.grid(row=1,column=1)

    edit_username_label = ttk.Label(edit_frame,text='username')
    edit_password_label = ttk.Label(edit_frame,text='new password')
    edit_username_label.grid(row=0,column=0)
    edit_password_label.grid(row=1,column=0)

    delete_button = ttk.Button(delete_frame,text='DELETE',command=edit)
    delete_button.grid(row=4,column=0,columnspan=2)

    delete_username_entry = ttk.Entry(delete_frame,width=50)
    delete_username_entry.grid(row=0,column=1)
    delete_username_label = ttk.Label(delete_frame,text='new password')
    delete_username_label.grid(row=0,column=0)



    



signin_frame = ttk.LabelFrame(mainframe,text='Al Mansoor Garments Signin',relief=SUNKEN)
signin_frame.grid(row=0,column=0,ipadx=10,ipady=10, padx=200,pady=70)

imglabel = Label(signin_frame,image=img)
imglabel.grid(row=0,column=0, columnspan=2)

username_label = ttk.Label(signin_frame, text='Username')
username_label.grid(row=1,column=0,pady=10)  

password_label = ttk.Label(signin_frame, text='Password')
password_label.grid(row=2,column=0,padx=10)

username =  ttk.Entry(signin_frame,width=50)    
username.grid(row=1,column=1,pady=10)    

password =  ttk.Entry(signin_frame,width=50, show='*')  
password.grid(row=2,column=1,pady=10)    

def sign():
    conn = sqlite3.connect('garment.db')
    c = conn.cursor()
    data = c.execute('SELECT * FROM auth')

    panel = ttk.LabelFrame(mainframe, text='Admin panel')
    panel.grid(row=1,column=0)

    for x in data:
        if username.get() == x[0] and password.get() == x[1] and username.get() == 'Shahid':
            signin_frame.grid_forget()

            hello = ttk.Label(panel,text=f'welcome {username.get()} you have sucessfully logged in')
            hello.grid(row=0,column=0)
            add = ttk.Button(panel, text='Add stock',command=add_stock)
            add.grid(row=5,column=1, sticky=W+E)

            sell_button = ttk.Button(panel, text='Sell',command=sell_stock)
            sell_button.grid(row=6,column=1, sticky=W+E)

            return_button = ttk.Button(panel, text='Return',command=return_stock)
            return_button.grid(row=7,column=1, sticky=W+E)

            show_detail = ttk.Button(panel, text='Detail', command=detail)
            show_detail.grid(row=8,column=1, sticky=W+E)

            user = ttk.Button(panel,text='user management', command=user_management)
            user.grid(row=9,column=1, sticky=W+E)

            daily = ttk.Button(panel,text="daily details", command=daily_detail)
            daily.grid(row=10,column=1, sticky=W+E)

        if username.get() == x[0] and password.get() == x[1] and username.get() != 'Shahid':
            signin_frame.grid_forget()

            hello = ttk.Label(panel,text=f'welcome {username.get()} you have sucessfully logged in')
            hello.grid(row=0,column=0, sticky=W+E)

            sell_button = ttk.Button(panel, text='Sell',command=sell_stock)
            sell_button.grid(row=6,column=1, sticky=W+E)

            return_button = ttk.Button(panel, text='Return',command=return_stock)
            return_button.grid(row=7,column=1, sticky=W+E)

        else:
            error = ttk.Label(signin_frame,text='Could not signin please enter valid credentials')
            error.grid(row=10,column=0,columnspan=20)


signin = ttk.Button(signin_frame, text='Sign in',command=sign)
signin.grid(row=3,column=0,columnspan=2) 
signin.bind('<Return>',sign)

root.mainloop()
    


